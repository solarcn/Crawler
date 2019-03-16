#coding:utf-8
import requests

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from PIL import Image
import requests
from io import BytesIO
import time
#from src.stringConvert import convertQ2B


#pattern = re.compile(r'www.{1,20}com.?】')
# patternBQG = re.compile(r'笔趣阁')

#笔趣阁 biquke.com
#八一中文 81xsw.com


def get_encoding(soup):
    encod = soup.meta.get('charset')
    if encod == None:
        encod = soup.meta.get('content-type')
        if encod == None:
            content = soup.meta.get('content')
            match = re.search('charset=(.*)', content)
            if match:
                encod = match.group(1)
                print("get_encoding()得到" + encod)
            else:
                raise ValueError('unable to find encoding')
    return encod

url1 = 'http://bbs.tianya.cn/post-stocks-841558-' 
page = 1
url2 = '.shtml'
lastpage = 202

document = Document()

fileName= '股王.docx'


 

def getContent(url):
    result = ''
    print(url)
    header = {'Referer':'%s' % url}
    r = requests.get(url)
#     print(r.headers.getparam('charset'))
#    r.encoding = 'gbk'
#     r.encoding = 'utf-8'
    soup = BeautifulSoup(r.text,'lxml')
    encoding = get_encoding(soup)
    if(encoding != None) :
        r.encoding = encoding;
        soup = BeautifulSoup(r.text,'lxml')

    records = soup.find_all('div', 'atl-item', js_username="股王01")



    for r in records:
        time = r.find_all('span')[1].text[3:]
        text = r.find_all('div', 'bbs-content')[0].text
        document.add_paragraph(time.strip())
        document.add_paragraph(text.strip())
        if(r.find_all('img') != []) :
            img = r.find_all('img')[0].get('original')
            response = requests.get(img,headers=header)
            imgIO = BytesIO(response.content)
            document.add_picture(imgIO, width=Inches(6.0))

 
#    novelContent = soup.find('div', id='content').contents

#    for line in novelContent:
#        string = convertQ2B(str(line)).lower()
##         print(string)
#        if string == '<br/>':
#            pass
##             result += '\n'
#        elif "】" in string:
#            print("原始:  " + string)
#            string = re.sub(pattern,'',string)
##             string = string.lstrip()[16:];
#            print("变换:  " + string)
#            result = result + string + '\n'
##             print("删除:" + string)
##         elif "笔趣阁" in string:
##             print("原始:  " + string)
##             string = string[:-4];
##             print("变换:  " + string)
##             result = result + string + '\n'
#        else:
#            result = result + string + '\n'
    return result, ""

#file = open(imageFileName, 'w', encoding='utf-8')

# i = 0; 
# while i < 678:
while page <= lastpage:
    getContent(url1 + str(page) + url2)
    document.add_paragraph('---------------------------------------------------------------------------------------------------EndOfPage' + str(page))
    page += 1 
#    time.sleep(0.1)
    
document.save(fileName)